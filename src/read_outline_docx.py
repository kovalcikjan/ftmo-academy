"""Read lesson outline .docx and output as plain text.

Extracts headings, paragraphs, tables, and bullet lists from the
outline .docx file. Used when resuming from Step 3 after a trading
expert has edited the outline.

Usage:
    python src/read_outline_docx.py data/output/lessons/[slug]/lesson_[slug]_EN_outline.docx
"""

import sys
from pathlib import Path

from docx import Document


def extract_text(docx_path: str) -> str:
    """Extract structured text from a .docx outline file.

    Args:
        docx_path: Path to the .docx file.

    Returns:
        Plain text representation of the document content.
    """
    doc_path = Path(docx_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")

    doc = Document(str(doc_path))
    lines: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            style_name = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            if style_name.startswith("Heading"):
                level = int(style_name.replace("Heading ", ""))
                prefix = "#" * level
                lines.append(f"{prefix} {text}")
            elif style_name == "List Bullet":
                lines.append(f"- {text}")
            elif style_name == "List Number":
                lines.append(f"1. {text}")
            else:
                lines.append(text)

        elif tag == "tbl":
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table is None:
                continue

            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python src/read_outline_docx.py <outline.docx>")
        sys.exit(1)

    output = extract_text(sys.argv[1])
    print(output)
