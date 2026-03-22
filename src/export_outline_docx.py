"""Export lesson outline/brief from markdown to .docx format.

Parses the outline .md file and generates a formatted Word document
with headings, tables, bullet lists, callouts, and content validation.

Usage:
    python src/export_outline_docx.py data/output/lessons/[slug]/lesson_[slug]_EN_outline.md
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# FTMO brand colors
FTMO_BLUE = RGBColor(0x12, 0x3A, 0x83)  # torea-bay
FTMO_AZURE = RGBColor(0x07, 0x81, 0xFE)  # azure
FTMO_GRAY = RGBColor(0x55, 0x55, 0x55)  # muted
FTMO_BODY = RGBColor(0x1A, 0x1A, 0x2E)  # body text


def create_document() -> Document:
    """Create a new document with FTMO Academy styling."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = FTMO_BODY

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = FTMO_BLUE
        heading_style.font.name = "Calibri"
        if level == 1:
            heading_style.font.size = Pt(22)
        elif level == 2:
            heading_style.font.size = Pt(16)
        elif level == 3:
            heading_style.font.size = Pt(13)
        else:
            heading_style.font.size = Pt(11)

    return doc


def add_table(doc: Document, lines: list[str]) -> None:
    """Parse markdown table lines and add a formatted table to the document."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells and all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        if cells:
            rows.append(cells)

    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j >= num_cols:
                break
            cell = row.cells[j]
            cell.text = cell_text
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["Normal"]

            if i == 0:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "123A83",
                    },
                )
                shading.append(shading_elem)

    doc.add_paragraph()


def add_callout(doc: Document, callout_type: str, text: str) -> None:
    """Add a styled callout paragraph."""
    labels = {
        "warning": "Warning",
        "tip": "Tip",
        "important": "Important",
        "note": "Note",
        "keep in mind": "Keep in Mind",
    }
    label = labels.get(callout_type.lower(), callout_type.title())

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.color.rgb = FTMO_AZURE

    run_text = p.add_run(text)
    run_text.font.color.rgb = FTMO_BODY


def parse_and_export(md_path: str) -> str:
    """Parse markdown outline and export as .docx.

    Args:
        md_path: Path to the outline .md file.

    Returns:
        Path to the generated .docx file.
    """
    md_file = Path(md_path)
    if not md_file.exists():
        raise FileNotFoundError(f"File not found: {md_path}")

    content = md_file.read_text(encoding="utf-8")
    lines = content.split("\n")

    doc = create_document()

    i = 0
    table_buffer: list[str] = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Flush table buffer if we leave a table
        if in_table and not stripped.startswith("|"):
            add_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        # Table row
        if stripped.startswith("|") and "|" in stripped[1:]:
            in_table = True
            table_buffer.append(stripped)
            i += 1
            continue

        # Horizontal rule — skip
        if stripped == "---":
            i += 1
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = re.sub(r"\s*\(H1\)\s*$", "", stripped[2:])
            doc.add_heading(title, level=1)
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            title = re.sub(r"\s*\(H2\)\s*$", "", stripped[3:])
            doc.add_heading(title, level=2)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            title = re.sub(r"\s*\(H3\)\s*$", "", stripped[4:])
            doc.add_heading(title, level=3)
            i += 1
            continue

        # H4
        if stripped.startswith("#### "):
            title = re.sub(r"\s*\(H4\)\s*$", "", stripped[5:])
            doc.add_heading(title, level=4)
            i += 1
            continue

        # Callout lines
        callout_match = re.match(
            r"^Callout:\s*(Warning|Tip|Important|Note|Keep in Mind)\s*[—–-]\s*(.+)$",
            stripped,
            re.IGNORECASE,
        )
        if callout_match:
            add_callout(doc, callout_match.group(1), callout_match.group(2))
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold at start
            bold_match = re.match(r"^\*\*(.+?)\*\*(.*)$", text)
            if bold_match:
                run = p.add_run(bold_match.group(1))
                run.bold = True
                p.add_run(bold_match.group(2))
            else:
                p.add_run(text)
            i += 1
            continue

        # Key point line
        if stripped.startswith("Key point:") or stripped.startswith("Key points:"):
            p = doc.add_paragraph()
            label_text = "Key point:" if stripped.startswith("Key point:") else "Key points:"
            rest = stripped[len(label_text):].strip()
            run = p.add_run(f"{label_text} ")
            run.bold = True
            run.font.color.rgb = FTMO_AZURE
            p.add_run(rest)
            i += 1
            continue

        # Target keyword line
        if stripped.startswith("Target keyword:"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run("Target keyword: ")
            run.bold = True
            run.font.color.rgb = FTMO_GRAY
            run.font.size = Pt(10)
            kw_run = p.add_run(stripped[len("Target keyword:"):].strip())
            kw_run.font.color.rgb = FTMO_GRAY
            kw_run.font.size = Pt(10)
            kw_run.italic = True
            i += 1
            continue

        # Regular text (non-empty)
        if stripped:
            # Handle bold fragments
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.+?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        # Empty line — skip
        i += 1

    # Flush remaining table
    if table_buffer:
        add_table(doc, table_buffer)

    # Save
    docx_path = md_file.with_suffix(".docx")
    doc.save(str(docx_path))
    return str(docx_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python src/export_outline_docx.py <outline.md>")
        sys.exit(1)

    output = parse_and_export(sys.argv[1])
    print(f"Generated: {output}")
